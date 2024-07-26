import os
from PyPDF2 import PdfReader
from docx import Document
from pdf2image import convert_from_path
from PIL import Image, ImageDraw, ImageFont

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        reader = PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text

def extract_text_from_word(docx_path):
    doc = Document(docx_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

def generate_image_from_pdf(pdf_path, output_image_path):
    images = convert_from_path(pdf_path, first_page=1, last_page=1)
    if images:
        images[0].save(output_image_path, 'PNG')

def generate_image_from_word(docx_path, output_image_path):
    doc = Document(docx_path)
    if doc.paragraphs:
        first_page_text = '\n'.join([para.text for para in doc.paragraphs[:20]])  # Limita el contenido de la primera página
        img = Image.new('RGB', (800, 1000), color=(255, 255, 255))
        d = ImageDraw.Draw(img)
        font = ImageFont.load_default()
        d.text((10, 10), first_page_text, fill=(0, 0, 0), font=font)
        img.save(output_image_path)

def extract_text_and_generate_images(directory):
    for filename in os.listdir(directory):
        if filename.endswith('.pdf'):
            pdf_path = os.path.join(directory, filename)
            print(f'Text from {filename}:\n{extract_text_from_pdf(pdf_path)}\n')
            generate_image_from_pdf(pdf_path, os.path.join(directory, f'{filename[:-4]}.png'))
        elif filename.endswith('.docx'):
            docx_path = os.path.join(directory, filename)
            print(f'Text from {filename}:\n{extract_text_from_word(docx_path)}\n')
            generate_image_from_word(docx_path, os.path.join(directory, f'{filename[:-5]}.png'))

# Asume que los archivos están en el mismo directorio que el script
current_directory = os.path.dirname(os.path.abspath(__file__))
extract_text_and_generate_images(current_directory)
